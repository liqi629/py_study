import pandas as pd

from docx import Document
from collections import Counter
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
#  excel配置文件
'''
一级模块使用2级标题，二级模块使用3级，三级模块使用4级
首先判断一级模块，添加一级模块和正文1，然后判断二级模块，添加二级模块和正文2
'''
level_1_menu=0
body_1=1
level_2_menu=2
body_2=3
level_3_menu=4
cast_tag = 5
case_num = 6
# 标志位，是否为用例的最后一个序号。
syc =7
# case_name =4
case_from = 8
case_illustrate =9
case_init =10
case_premise =11
case_step = 13
# 预期
expected_result = 14
# 评估
evaluative_criteria = 15
# 备注
mark = 16
# 测试设计
tester =17
date=18


df = pd.read_excel('E:\YOYO\测试用例\echo_gjb_case.xlsx')
data = df.values
# print(data)
# 用例标识集合

# 创建一个word，添加一个表格
document = Document()
table = None
last_tag = None

# 设置一个空白样式
style = document.styles['Normal']
# 设置中文字体
# style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for i in range(len(data)):
    # print(data[i])
    row_data = data[i]
    case_tag_d = row_data[cast_tag]
    body_1_d = row_data[body_1]
    body_2_d = row_data[body_2]

    if body_1_d !='1':
        document.add_heading(str(row_data[level_1_menu]) , 2)

        paragraph = document.add_paragraph(str(row_data[body_1]))

        # paragraph_format = style.paragraph_format
        # paragraph_format.first_line_indent = Cm(0.74)

    if body_2_d !='1' and body_2_d !='None'and body_2_d !='skip':
        document.add_heading(str(row_data[level_2_menu]), 3)
        paragraph = document.add_paragraph(str(row_data[body_2]))
    elif body_2_d =='None':
        document.add_heading(str(row_data[level_2_menu]), 3)

    # 如果新的测试用例标识和上一个而不一样，添加标题，和部分表格内容
    if last_tag != case_tag_d:
        # 判断一下，是不是有二级标题，否则将三级标题提高一级
        if body_2_d !='skip':
            document.add_heading(str(row_data[level_3_menu])+"("+str(row_data[cast_tag]) + ') ' , 4)
        elif body_2_d =='skip':
            document.add_heading(str(row_data[level_3_menu]) + "(" + str(row_data[cast_tag]) + ') ', 3)
        table = document.add_table(rows=7, cols=5, style='Table Grid')
        # print("这条case，与之前的都不属于一个功能模块，需要新建一个表格写入\n测试用例的标识是：{0}，序号是{1}".format(row_data[cast_tag], row_data[case_num]))

        table.cell(0, 0).text = u'用例名称'
        table.cell(0, 1).text = str(row_data[level_3_menu])

        table.cell(0, 2).text = u'用例标识'
        table.cell(0, 3).merge(table.cell(0,4))
        table.cell(0, 3).text = str(row_data[cast_tag])

        table.cell(1, 0).text = u'测试追踪'
        table.cell(1, 1).merge(table.cell(1, 4))
        table.cell(1, 1).text = str(row_data[case_from])

        table.cell(2, 0).text = u'测试说明'
        table.cell(2, 1).merge(table.cell(2, 4))
        table.cell(2, 1).text = str(row_data[case_illustrate])

        table.cell(3, 0).text = u'用例初始化'
        table.cell(3, 1).merge(table.cell(3, 4))
        table.cell(3, 1).text = str(row_data[case_init])

        table.cell(4, 0).merge(table.cell(4, 4))

        table.cell(4, 0).text = u'测试过程'
        table.cell(4, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER


        table.cell(5, 0).text = u'前提约束'
        table.cell(5, 1).merge(table.cell(5, 4))
        table.cell(5, 1).text = str(row_data[case_premise])

        table.cell(6, 0).text = u'序号'
        table.cell(6, 1).text = u'输入及操作说明'
        table.cell(6, 2).text = u'期望'
        table.cell(6, 3).text = u'评估标准'
        table.cell(6, 4).text = u'备注'

        table.add_row()
        table.cell(-1, 0).text = str(row_data[case_num])
        table.cell(-1, 1).text = str(row_data[case_step])
        table.cell(-1, 2).text = str(row_data[expected_result])
        table.cell(-1, 3).text = str(row_data[evaluative_criteria])
        table.cell(-1, 4).text = str(row_data[mark])
        # 如果标志位为2，补全表格剩余部分
        if row_data[syc] == 2:
            table.add_row()
            table.cell(-1, 1).merge(table.cell(-1, 4))
            table.cell(-1, 0).text = u'测试终止条件'

            table.cell(-1, 1).text = u'测试结果符合预期后，测试终止'
            table.add_row()
            table.cell(-1, 1).merge(table.cell(-1, 4))
            table.cell(-1, 0).text = u'测试结果'
            table.cell(-1, 1).text = u'□通过  □未通过但可作优化或修改  □未通过且无法修改'
            table.add_row()
            table.cell(-1, 0).text = u'设计人员'
            table.cell(-1, 1).text = str(row_data[tester])

            table.cell(-1, 2).text = u'设计日期'
            table.cell(-1, 3).merge(table.cell(-1, 4))
            table.cell(-1, 3).text = str(row_data[date])
        last_tag = case_tag_d
        table = table
    # 如果 测试用例标识与上一个一样
    else:
        # 判断标志位是不是不是2，继续添加表格内容
        if row_data[syc] !=2:
            table.add_row()
            table.cell(-1, 0).text = str(row_data[case_num])
            table.cell(-1, 1).text = str(row_data[case_step])
            table.cell(-1, 2).text = str(row_data[expected_result])
            table.cell(-1, 3).text = str(row_data[evaluative_criteria])
            table.cell(-1, 4).text = str(row_data[mark])
        # 如果标志位是2，添加表格内容后，补全表格剩余部分
        else:
            table.add_row()
            table.cell(-1, 0).text = str(row_data[case_num])
            table.cell(-1, 1).text = str(row_data[case_step])
            table.cell(-1, 2).text = str(row_data[expected_result])
            table.cell(-1, 3).text = str(row_data[evaluative_criteria])
            table.cell(-1, 4).text = str(row_data[mark])

            table.add_row()
            table.cell(-1, 1).merge(table.cell(-1, 4))
            table.cell(-1, 0).text = u'测试终止条件'

            table.cell(-1, 1).text = u'测试结果符合预期后，测试终止'
            table.add_row()
            table.cell(-1, 1).merge(table.cell(-1, 4))
            table.cell(-1, 0).text = u'测试结果'
            table.cell(-1, 1).text = u'□通过  □未通过但可作优化或修改  □未通过且无法修改'
            table.add_row()
            table.cell(-1, 0).text = u'设计人员'
            table.cell(-1, 1).text = str(row_data[tester])

            table.cell(-1, 2).text = u'设计日期'
            table.cell(-1, 3).merge(table.cell(-1, 4))
            table.cell(-1, 3).text = str(row_data[date])


#
document.add_page_break()
document.save('23.docx')



